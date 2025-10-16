# fill_card.py
import pandas as pd
from docx import Document
import re
import os


def replace_in_paragraph(para, replacements):
    full_text = "".join(run.text for run in para.runs)
    if not any(ph in full_text for ph in replacements):
        return

    for placeholder, value in replacements.items():
        full_text = full_text.replace(placeholder, value)

    if para.runs:
        para.runs[0].text = full_text
        for run in para.runs[1:]:
            para._element.remove(run._element)

def format_address_nicely(address: str) -> str:
    if not address.strip():
        return address

    result = address  # работаем с оригиналом

    # 🔥 1. Москва — ваша рабочая логика
    if 'Г.МОСКВА' in result.upper() or 'Г. МОСКВА' in result.upper():
        result = re.sub(
            r'(г\.\s*Москва,\s*)МУНИЦИПАЛЬНЫЙ ОКРУГ АЭРОПОРТ',
            r'г. Москва, вн.тер. Муниципальный округ Аэропорт',
            result,
            flags=re.IGNORECASE
        )

    # 🔥 2. Сириус — ваша рабочая логика
    elif 'СИРИУС' in result.upper() and ('КРАЙ' in result.upper() or 'ОБЛАСТЬ' in result.upper()):
        result = re.sub(
            r'(КРАСНОДАРСКИЙ КРАЙ),\s*СИРИУС.*?(ПРОЕЗД|УЛ\.|УЛИЦА)',
            r'Краснодарский край, ф.т. Сириус, пгт. Сириус, \2',
            result,
            flags=re.IGNORECASE
        )

    # 🔥 3. Иннополис — ПО АНАЛОГИИ с Москвой и Сириусом
    elif 'РЕСПУБЛИКА ТАТАРСТАН' in result.upper() and 'ИННОПОЛИС' in result.upper():
        result = re.sub(
            r'(РЕСПУБЛИКА ТАТАРСТАН\s*\(ТАТАРСТАН\)),\s*ВЕРХНЕУСЛОНСКИЙ,\s*Г\s+ИННОПОЛИС,\s*(УЛ\.?|УЛИЦА)',
            r'Республика Татарстан (Татарстан), Верхнеуслонский м.р-н, г.п. город Иннополис, г. Иннополис, ул.',
            result,
            flags=re.IGNORECASE
        )

    # 🔥 4. Общая стандартизация — ТОЛЬКО сокращения с точкой
    result = re.sub(r'\bПОМЕЩ\.?\s*', 'помещ. ', result, flags=re.IGNORECASE)
    result = re.sub(r'\bПРОЕЗД\.?\s*', 'проезд ', result, flags=re.IGNORECASE)
    result = re.sub(r'\bРАБ. МЕСТО\.?\s*', 'раб. место ', result, flags=re.IGNORECASE)
    result = re.sub(r'\bД\.?\s*', 'д. ', result, flags=re.IGNORECASE)
    result = re.sub(r'\bУЛ\.?\s*', 'ул. ', result, flags=re.IGNORECASE)
    result = re.sub(r'\bПР-КТ\.?\s*', 'пр-кт ', result, flags=re.IGNORECASE)
    result = re.sub(r'\bГ\.\s*', 'г. ', result, flags=re.IGNORECASE)  # только "Г.", не "Г "

    # Капитализация улиц
    result = re.sub(
        r'(ул\.|пр-кт|ш\.|проезд)\s+([а-яё-]+)',
        lambda m: m.group(1) + " " + m.group(2).capitalize(),
        result,
        flags=re.IGNORECASE
    )

    return re.sub(r'\s+', ' ', result).strip()

    print(f"📤 Результат: '{result}'")
    return result

def fill_card_for_company(df, template_path, output_dir, full_company_name):
    match = df[df["НаимПолн"] == full_company_name]
    if match.empty:
        print(f"⚠️ Компания '{full_company_name}' не найдена.")
        return False

    row = match.iloc[0].to_dict()
    print("\n" + "=" * 60)
    print(f"🔍 Обработка компании: {full_company_name}")
    print(f"📊 Столбцы в Excel: {list(row.keys())}")
    raw_address = str(row.get("Адрес", "ОТСУТСТВУЕТ"))
    print(f"📥 Адрес из Excel: '{raw_address}'")

    formatted_address = format_address_nicely(raw_address)
    print(f"📤 Адрес после обработки: '{formatted_address}'")

    ogrn_val = str(row.get("ОГРН", "")).strip()
    inn_val = str(row.get("ИНН", "")).strip()
    kpp_val = str(row.get("КПП", "")).strip()
    address_val = formatted_address  # ← ВАЖНО: используем ОБРАБОТАННЫЙ адрес!
    company_val = str(row.get("НаимПолн", "")).strip().title()
    title_val = str(row.get("НаимСокр", "")).strip()
    print(f"🔍 Сокращённое наименование: '{title_val}'")
    okved_val = str(row.get("ОКВЭД_Основной", "")).strip()
    okved1_val = str(row.get("ОКВЭД_Доп", "")).strip()
    mail_val = str(row.get("Email", "")).strip().title()
    post_val = str(row.get("Должность", "")).strip().title()
    fio_val = str(row.get("ФИО_Руководителя", "")).strip().title()

#Создаем словарь из значений

    replacements = {
        "{{ogrn}}": ogrn_val,
        "{{inn}}": inn_val,
        "{{kpp}}": kpp_val,
        "{{address}}": address_val,
        "{{company}}": company_val,
        "{{title}}": title_val,
        "{{okved}}": okved_val,
        "{{okved1}}": okved1_val,
        "{{mail}}": mail_val,
        "{{post}}": post_val,
        "{{fio}}": fio_val,
        "{{ ogrn }}": ogrn_val,
        "{{ inn }}": inn_val,
        "{{ kpp }}": kpp_val,
        "{{ address }}": address_val,
        "{{ company }}": company_val,
        "{{ title }}": title_val,
        "{{ okved }}": okved_val,
        "{{ okved1 }}": okved1_val,
        "{{ mail }}": mail_val,
        "{{ post }}": post_val,
        "{{ fio }}": fio_val,
    }

    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"❌ Не удалось открыть шаблон: {e}")
        return False

    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)

    safe_name = re.sub(r'[<>:"/\\|?*]', '_', full_company_name)
    output_path = os.path.join(output_dir, f"{safe_name}.docx")

    try:
        doc.save(output_path)
        print(f"✅ Сохранено: {output_path}")
        return True
    except PermissionError:
        print(f"❌ Ошибка: файл '{output_path}' открыт в Word. Закройте его.")
        return False
    except Exception as e:
        print(f"❌ Ошибка сохранения: {e}")
        return False


def main():
    EXCEL_FILE = "egrul_multiple.xlsx"
    TEMPLATE_FILE = "C:/Users/Екатерина/PycharmProjects/HSE_Kate_Shadr/EGRUL//Карточка Т1.docx"
    OUTPUT_DIR = "C:/Users/Екатерина/PycharmProjects/HSE_Kate_Shadr/EGRUL/Актуализация/заполненные_карточки"

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    COMPANY_NAMES = [
        'ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ "НОТА СЕРВИС"',
        'АКЦИОНЕРНОЕ ОБЩЕСТВО "Т1"',
        'АВТОНОМНАЯ НЕКОММЕРЧЕСКАЯ ОРГАНИЗАЦИЯ ДОПОЛНИТЕЛЬНОГО ПРОФЕССИОНАЛЬНОГО ОБРАЗОВАНИЯ "Т1 ЦИФРОВАЯ АКАДЕМИЯ"'
    ]

    try:
        df = pd.read_excel(EXCEL_FILE)
    except Exception as e:
        print(f"❌ Не удалось прочитать Excel: {e}")
        return

    success_count = 0
    for company in COMPANY_NAMES:
        print(f"\nОбработка: {company}")
        if fill_card_for_company(df, TEMPLATE_FILE, OUTPUT_DIR, company):
            success_count += 1

    print(f"\n✅ Успешно заполнено документов: {success_count} из {len(COMPANY_NAMES)}")


if __name__ == "__main__":
    main()