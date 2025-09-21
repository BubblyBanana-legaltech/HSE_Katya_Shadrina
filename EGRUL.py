import requests
import xml.etree.ElementTree as ET
import pandas as pd
from datetime import datetime
import os


def parse_org_xml(xml_content):
    root = ET.fromstring(xml_content)
    sv_yul = root.find("СвЮЛ")
    if sv_yul is not None:
        attrs = sv_yul.attrib
        sv_naim_yul = sv_yul.find("СвНаимЮЛ/СвНаимЮЛСокр")
        name_short = sv_naim_yul.attrib["НаимСокр"] if sv_naim_yul is not None else ""
        sv_naim_yul_poln = sv_yul.find("СвНаимЮЛ/СвНаимЮЛПолн")
        name_full = sv_naim_yul_poln.attrib["НаимЮЛПолн"] if sv_naim_yul_poln is not None else ""
        sv_adres_yul = sv_yul.find("СвАдресЮЛ/АдресРФ")
        address = ""
        if sv_adres_yul is not None:
            region = sv_adres_yul.find("Регион")
            street = sv_adres_yul.find("Улица")
            house = sv_adres_yul.find("Дом")
            address_parts = []
            if region is not None and "НаимРегион" in region.attrib:
                address_parts.append(region.attrib["НаимРегион"])
            if street is not None and "НаимУлица" in street.attrib:
                address_parts.append(street.attrib["НаимУлица"])
            if house is not None and house.text:
                address_parts.append(house.text)
            address = ", ".join(address_parts)
        return {
            "ОГРН": attrs.get("ОГРН", ""),
            "ИНН": attrs.get("ИНН", ""),
            "НаимСокр": name_short,
            "НаимПолн": name_full,
            "Адрес": address,
            "ДатаВып": attrs.get("ДатаВып", ""),
            "ДатаЗагрузки": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
    return None


def get_org_data(ogrn):
    url = f"https://egrul.itsoft.ru/{ogrn}.xml"
    response = requests.get(url)
    if response.status_code == 200 and response.text != 'false':
        return parse_org_xml(response.content)
    else:
        print(f"Данные не найдены для ОГРН {ogrn}")
        return None


def save_to_excel(data_list, filename='egrul_multiple.xlsx'):
    if not data_list:
        print("Нет данных для сохранения")
        return

    if os.path.exists(filename):
        # Если файл существует, считываем существующий DataFrame и добавляем новые данные
        existing_df = pd.read_excel(filename)
        new_df = pd.DataFrame(data_list)
        combined_df = pd.concat([existing_df, new_df], ignore_index=True)
        combined_df.to_excel(filename, index=False)
    else:
        # Если файла нет, просто создаём новый
        df = pd.DataFrame(data_list)
        df.to_excel(filename, index=False)
    print(f"Данные успешно сохранены в {filename}")


if __name__ == "__main__":
    ogrn_list = ["1227700086460", "1027700132195", "1257700213847"]  # Пример списка ОГРН
    all_data = []
    for ogrn in ogrn_list:
        data = get_org_data(ogrn)
        if data:
            all_data.append(data)
    save_to_excel(all_data)

    import pandas as pd
    from docx import Document

    # Читаем Excel
    df = pd.read_excel('egrul_multiple.xlsx')

    # Создаем Word документ
    doc = Document()
    doc.add_heading('Данные из Excel', level=1)

    # Создаем таблицу в Word: строки + 1 для заголовков, столбцы = столбцам Excel
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

    # Заполняем заголовки таблицы
    for j, col_name in enumerate(df.columns):
        table.cell(0, j).text = str(col_name)

    # Заполняем данные
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            val = df.iat[i, j]
            table.cell(i + 1, j).text = "" if pd.isna(val) else str(val)

    # Сохраняем Word файл
    doc.save('output.docx')
    print("Word документ создан успешно")